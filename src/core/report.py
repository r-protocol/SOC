# report.py
import sqlite3
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.config import DATABASE_PATH, TEMPLATE_DOCX_PATH, OUTPUT_DOCX_PATH
from src.utils.logging_utils import log_info, log_success, log_warn, log_error

def get_last_full_week_dates():
    import datetime
    today = datetime.date.today()
    end_date = today - datetime.timedelta(days=(today.weekday() + 1))
    start_date = end_date - datetime.timedelta(days=6)
    return start_date, end_date

def add_table_borders(table):
    tbl_pr = table._element.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_elm = OxmlElement(f'w:{border_name}')
        border_elm.set(ns.qn('w:val'), 'single')
        border_elm.set(ns.qn('w:sz'), '4')
        border_elm.set(ns.qn('w:space'), '0')
        border_elm.set(ns.qn('w:color'), 'auto')
        tbl_borders.append(border_elm)
    tbl_pr.append(tbl_borders)

def set_cell_background_color(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), color_hex.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_bookmark(paragraph, bookmark_name):
    run = paragraph.runs[0]
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(ns.qn('w:id'), '0')
    start.set(ns.qn('w:name'), bookmark_name)
    tag.addprevious(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(ns.qn('w:id'), '0')
    tag.addnext(end)

def add_internal_hyperlink(paragraph, text, anchor):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(ns.qn('w:anchor'), anchor)
    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    style = OxmlElement('w:rStyle')
    style.set(ns.qn('w:val'), 'Hyperlink')
    run_props.append(style)
    run.append(run_props)
    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    for r in paragraph.runs:
        if r.text == text:
            r.font.name = 'Arial'
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.font.underline = True

def generate_weekly_report():
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    start_date, end_date = get_last_full_week_dates()
    cursor.execute("SELECT * FROM articles WHERE published_date BETWEEN ? AND ?", (start_date.isoformat(), end_date.isoformat()))
    all_weekly_articles = cursor.fetchall()
    conn.close()
    if not all_weekly_articles:
        log_warn("No articles found from the last week. No report generated.")
        return
    log_info(f"Found {len(all_weekly_articles)} total articles in DB for the week. Selecting based on risk quota...")
    categorized_articles = {"HIGH": [], "MEDIUM": [], "LOW": [], "INFORMATIONAL": []}
    for article in all_weekly_articles:
        risk = (article[6] or "N/A").upper()
        if risk in categorized_articles:
            categorized_articles[risk].append(article)
    report_quota = {"HIGH": 4, "MEDIUM": 3, "LOW": 2, "INFORMATIONAL": 1}
    target_report_size = 10
    weekly_data = []
    for risk_level in ["HIGH", "MEDIUM", "LOW", "INFORMATIONAL"]:
        count = report_quota[risk_level]
        available = categorized_articles.get(risk_level, [])
        selected = available[:count]
        weekly_data.extend(selected)
        for item in selected:
            categorized_articles[risk_level].remove(item)
        if len(selected) < count:
            log_warn(f"Requirement not met for {risk_level}: Found {len(selected)}, need {count}.")
    if len(weekly_data) < target_report_size:
        log_info(f"Report has {len(weekly_data)} articles, filling to {target_report_size} with best available.")
        remaining_pool = []
        for risk_level in ["HIGH", "MEDIUM", "LOW", "INFORMATIONAL"]:
            remaining_pool.extend(categorized_articles[risk_level])
        needed = target_report_size - len(weekly_data)
        weekly_data.extend(remaining_pool[:needed])
    risk_priority = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "INFORMATIONAL": 3}
    weekly_data.sort(key=lambda row: risk_priority.get((row[6] or "N/A").upper(), 99))
    if not weekly_data:
        log_warn("No articles selected for the report after applying quotas.")
        return
    log_info(f"Generating report with {len(weekly_data)} selected articles.")
    try:
        doc = Document(TEMPLATE_DOCX_PATH)
    except Exception as e:
        log_error(f"Could not open '{TEMPLATE_DOCX_PATH}'. Error: {e}")
        return
    period_p, titles_p, summary_intro_p = None, None, None
    for p in doc.paragraphs:
        if '[Placeholder-ReportingPeriod]' in p.text: period_p = p
        if '[Placeholder-Titles]' in p.text: titles_p = p
        if 'This report provides a detailed summary' in p.text: summary_intro_p = p
    if period_p:
        period_p.text = f"Reporting period: {start_date.strftime('%d %b %Y')} – {end_date.strftime('%d %b %Y')}"
        for run in period_p.runs: run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Arial'
        period_p._p.addnext(OxmlElement('w:p'))
    if titles_p:
        parent_body = titles_p._element.getparent()
        placeholder_index = parent_body.index(titles_p._element)
        for i, row_data in enumerate(weekly_data):
            p = doc.add_paragraph(f"{i+1}. ")
            add_internal_hyperlink(p, row_data[1], f"threat_{i+1}")
            p.paragraph_format.left_indent = Inches(0.5)
            p_element = p._element
            p_element.getparent().remove(p_element)
            parent_body.insert(placeholder_index + i, p_element)
        parent_body.remove(titles_p._element)
    if summary_intro_p: summary_intro_p._p.addnext(OxmlElement('w:p'))
    for table in doc.tables: table._tbl.getparent().remove(table._tbl)
    risk_colors = {"HIGH": "980000", "MEDIUM": "FFC000", "LOW": "38761D", "INFORMATIONAL": "1155CC", "N/A": "A9A9A9"}
    for i, article_data in enumerate(weekly_data):
        _id, title, url, _, _, summary_text, threat_risk, _, recommendations_json = article_data
        summary_text = summary_text or "Summary not available."
        threat_risk = (threat_risk or "N/A").upper()
        recommendations = json.loads(recommendations_json) if recommendations_json else []
        doc.add_paragraph()
        article_table = doc.add_table(rows=0, cols=1)
        add_table_borders(article_table)
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"THREAT RISK = {threat_risk}"); run.font.name = 'Arial'; run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background_color(cell, risk_colors.get(threat_risk, "A9A9A9"))
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title); run.font.name = 'Arial'; run.font.size = Pt(14); run.font.bold = True
        add_bookmark(p, f"threat_{i+1}")
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SUMMARY"); run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True
        set_cell_background_color(cell, "CCCCCC")
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; run = p.add_run(summary_text); run.font.name = 'Arial'; run.font.size = Pt(11)
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RECOMMENDATIONS"); run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True
        set_cell_background_color(cell, "CCCCCC")
        cell = article_table.add_row().cells[0]
        p_to_remove = cell.paragraphs[0]._element
        p_to_remove.getparent().remove(p_to_remove)
        for reco in recommendations:
            p = cell.add_paragraph()
            if isinstance(reco, dict) and 'title' in reco and 'description' in reco:
                p.add_run("• ").font.bold = True; p.add_run(f"{reco['title']}: ").font.bold = True; p.add_run(reco['description'])
                for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)
            else:
                reco_text = reco.get('step', reco) if isinstance(reco, dict) else reco
                run = p.add_run(f"• {reco_text}"); run.font.name = 'Arial'; run.font.size = Pt(11)
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Sources"); run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True
        set_cell_background_color(cell, "CCCCCC")
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; run = p.add_run(f"• {url}"); run.font.name = 'Arial'; run.font.size = Pt(11)
    for paragraph in doc.paragraphs:
        p_format = paragraph.paragraph_format; p_format.line_spacing = 1.15; p_format.space_before = Pt(6); p_format.space_after = Pt(6)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    p_format = paragraph.paragraph_format; p_format.line_spacing = 1.15; p_format.space_before = Pt(6); p_format.space_after = Pt(6)
    try:
        doc.save(OUTPUT_DOCX_PATH)
        log_success(f"Word report successfully generated and saved to: {OUTPUT_DOCX_PATH}")
    except Exception as e:
        log_error(f"Could not save the document. Error: {e}")
