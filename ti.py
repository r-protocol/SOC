# This file has been split into multiple modules for better maintainability.
# Please use main.py as the entry point for the pipeline.
# Modules: config.py, logging_utils.py, db_utils.py, fetcher.py, filtering.py, analysis.py, report.py, main.py
# threat_intelligence_pipeline.py
#
# A modular pipeline to fetch, analyze, and report on cybersecurity threat intelligence.
# This version uses a sequential, per-phase approach for a clear and easy-to-follow output.

import feedparser
import requests
import json
import datetime
import random
import sqlite3
import re
import os
import sys
import socket
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm

# --- CONFIGURATION ---
RSS_FEEDS = [
    "https://www.bleepingcomputer.com/feed/",
    "https://feeds.feedburner.com/TheHackersNews",
    "https://www.darkreading.com/rss_simple.asp",
    "https://threatpost.com/feed/",
    "https://krebsonsecurity.com/feed/"
]
DATABASE_PATH = "threat_intel.db"
OLLAMA_MODEL = "llama3"
OLLAMA_HOST = "http://localhost:11434"
TEMPLATE_DOCX_PATH = "template.docx"
OUTPUT_DOCX_PATH = f"Threat_Intelligence_Report_{datetime.date.today()}.docx"
# --- END OF CONFIGURATION ---

# --- DEBUG MODE FLAG ---
DEBUG_MODE = False

# --- CONSOLE LOGGING SETUP ---
class BColors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

def log_info(message):
    print(f"{BColors.OKCYAN}[INFO]{BColors.ENDC} {message}")

def log_success(message):
    print(f"{BColors.OKGREEN}[SUCCESS]{BColors.ENDC} {message}")

def log_warn(message):
    print(f"{BColors.WARNING}[WARN]{BColors.ENDC} {message}")
    
def log_error(message):
    print(f"{BColors.FAIL}[ERROR]{BColors.ENDC} {message}")

def log_step(step_number, message):
    print(f"\n{BColors.BOLD}{BColors.HEADER}--- PHASE {step_number}: {message} ---{BColors.ENDC}")

USER_AGENTS = [
    'Mozilla/50.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/109.0.0.0 Safari/537.36',
    'Mozilla/50.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/109.0.0.0 Safari/537.36',
    'Mozilla/50.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Firefox/109.0 Safari/537.36',
]

def initialize_database():
    """Initializes the SQLite database for storing article data."""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS articles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            url TEXT UNIQUE NOT NULL,
            published_date TEXT NOT NULL,
            content TEXT,
            summary TEXT,
            threat_risk TEXT,
            category TEXT,
            recommendations TEXT
        )
    """)
    conn.commit()
    conn.close()
    log_success("Database initialized successfully.")

def get_existing_urls():
    """Fetches all unique URLs from the database to avoid re-processing."""
    if not os.path.exists(DATABASE_PATH):
        return set()
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT url FROM articles")
    urls = {row[0] for row in cursor.fetchall()}
    conn.close()
    return urls

def get_last_full_week_dates():
    """Calculates the start and end dates for the previous full calendar week (Monday to Sunday)."""
    today = datetime.date.today()
    end_date = today - datetime.timedelta(days=(today.weekday() + 1))
    start_date = end_date - datetime.timedelta(days=6)
    return start_date, end_date

# --- PHASE 1: FETCHING AND SCRAPING (Sequential) ---
def fetch_and_scrape_articles_sequential(existing_urls):
    """Fetches new articles from all RSS feeds sequentially."""
    all_articles = []
    start_date, end_date = get_last_full_week_dates()
    log_info(f"Searching for new articles from {len(RSS_FEEDS)} sources for the week of {start_date} to {end_date}...")

    for feed_url in tqdm(RSS_FEEDS, desc="Fetching from feeds"):
        try:
            socket.setdefaulttimeout(20)
            feed = feedparser.parse(feed_url)
            socket.setdefaulttimeout(None)
            
            for entry in feed.entries:
                if entry.link in existing_urls:
                    continue

                try:
                    published_date = datetime.date(*entry.published_parsed[:3]) if hasattr(entry, 'published_parsed') else datetime.date.today()
                except Exception:
                    published_date = datetime.date.today()

                if start_date <= published_date <= end_date:
                    print(f"\n{BColors.OKCYAN}[NEW]{BColors.ENDC} Fetched: {entry.title}")
                    article_data = {'title': entry.title, 'url': entry.link, 'published_date': published_date.isoformat(), 'content': ""}
                    if hasattr(entry, 'content') and entry.content:
                        article_data['content'] = BeautifulSoup(entry.content[0].value, 'html.parser').get_text(strip=True)
                    elif hasattr(entry, 'summary'):
                        article_data['content'] = BeautifulSoup(entry.summary, 'html.parser').get_text(strip=True)

                    if len(article_data['content']) < 200:
                        try:
                            headers = {'User-Agent': random.choice(USER_AGENTS)}
                            response = requests.get(article_data['url'], headers=headers, timeout=15)
                            response.raise_for_status()
                            soup = BeautifulSoup(response.content, 'html.parser')
                            body = soup.find('article') or soup.find('div', class_='article-body')
                            if body:
                                article_data['content'] = body.get_text(separator='\n', strip=True)
                        except requests.RequestException:
                            article_data['content'] = None
                    
                    all_articles.append(article_data)
        except (socket.timeout, Exception) as e:
            log_warn(f"Could not process feed {feed_url}: {e}")
            socket.setdefaulttimeout(None)
            
    log_success(f"Found a total of {len(all_articles)} new potential articles across all feeds.")
    return all_articles

# --- PHASE 2: RELEVANCE FILTERING (Sequential) ---
def filter_articles_sequential(articles):
    """Filters a list of articles for cybersecurity relevance sequentially."""
    relevant_articles = []
    articles_to_check = [a for a in articles if a.get('content')]
    
    if not articles_to_check:
        return []

    for article in tqdm(articles_to_check, desc="Filtering for relevance"):
        if is_article_relevant_with_llm(article):
            relevant_articles.append(article)
            print(f"\n{BColors.OKGREEN}[RELEVANT]{BColors.ENDC} {article['title']}")

    log_success(f"Found {len(relevant_articles)} new relevant articles to analyze.")
    return relevant_articles

def is_article_relevant_with_llm(article):
    """Uses the LLM to quickly classify if an article is cybersecurity-related."""
    prompt = f"""
You are a cybersecurity news curator. Your task is to determine if the following article is related to cybersecurity.
Cybersecurity topics include vulnerabilities, data breaches, malware, ransomware, phishing, threat actors, cyber attacks, and security advisories.
General technology news, such as product releases (e.g., new phones or gaming features), company financials, or general IT trends, are NOT relevant.

Based on the title and content, is this article about cybersecurity? Respond with only a single word: YES or NO.

Article Title: {article['title']}
Article Content: {article['content'][:500]}
"""
    try:
        response = requests.post(
            f"{OLLAMA_HOST}/api/generate",
            json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False},
            timeout=60
        )
        response.raise_for_status()
        response_text = response.json()['response'].strip().upper()
        return "YES" in response_text
    except requests.RequestException:
        return False

# --- PHASE 3: FULL ANALYSIS (Sequential) ---
def analyze_articles_sequential(articles):
    """Analyzes a list of articles using the LLM sequentially."""
    analyzed_articles = []
    for article in tqdm(articles, desc="Analyzing articles"):
        llm_analysis = analyze_article_with_llm(article)
        if llm_analysis:
            article.update(llm_analysis)
            analyzed_articles.append(article)
            print(f"\n{BColors.OKGREEN}[ANALYZED]{BColors.ENDC} {article['title']} (Risk: {article.get('threat_risk')})")
    return analyzed_articles

def repair_and_parse_json(raw_text):
    """Attempts to repair and parse a JSON object from a raw string that may contain errors."""
    json_start = raw_text.find('{')
    json_end = raw_text.rfind('}')
    if json_start == -1 or json_end == -1: return None
    json_str = raw_text[json_start:json_end + 1]
    json_str = re.sub(r'}\s*{', '},{', json_str)
    json_str = re.sub(r',\s*([}\]])', r'\1', json_str)
    
    repaired_str = ""
    in_string = False
    for char in json_str:
        if char == '"':
            in_string = not in_string
        if in_string and char == '\n':
            repaired_str += '\\n'
        else:
            repaired_str += char
    json_str = repaired_str

    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        return None

def analyze_article_with_llm(article):
    """Sends article content to an LLM for analysis, with retries."""
    prompt = f"""
You are a senior cybersecurity threat intelligence analyst. Your final output MUST be a single, valid, raw JSON object. Do not use markdown like ```json or **. Do not add any conversational text. The entire response must start with {{ and end with }}.

Analyze the following article and provide a threat assessment based on these strict definitions:
- **HIGH**: Active exploitation of a critical vulnerability, major data breach at a large company, or a widespread, severe malware campaign. Requires immediate action.
- **MEDIUM**: A significant vulnerability has been disclosed but isn't yet widely exploited, a smaller-scale breach, or a notable new malware variant. Requires timely review.
- **LOW**: A minor vulnerability, a theoretical attack vector, or a security advisory about a niche product. Should be monitored.
- **INFORMATIONAL**: General cybersecurity news, trend reports, or expert opinions that do not represent a direct threat. For awareness only.

1.  **summary**: A detailed, professional summary as a single JSON string, in a natural, narrative style (2-3 paragraphs, using "\\n\\n" for breaks). Explain what happened, why it's important, and the business impact.
2.  **threat_risk**: Must be one of: "HIGH", "MEDIUM", "LOW", "INFORMATIONAL".
3.  **category**: Must be one of: "Ransomware", "Phishing", "Vulnerability", "Malware", "Breach", "General Security".
4.  **recommendations**: A list of 5 actionable recommendations. Each must be a JSON object with two keys: a short "title" and a detailed "description".

Article Title: {article['title']}
Article Content: {article['content'][:8000]}
"""
    max_retries = 2
    for attempt in range(max_retries + 1):
        try:
            response = requests.post(
                f"{OLLAMA_HOST}/api/generate",
                json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False},
                timeout=300
            )
            response.raise_for_status()
            response_text = response.json()['response'].strip()
            parsed_json = repair_and_parse_json(response_text)
            if parsed_json:
                return parsed_json # Success
            elif attempt < max_retries:
                 print(f"\n{BColors.WARNING}[RETRYING]{BColors.ENDC} Invalid JSON for '{article['title']}' (Attempt {attempt + 2})")

        except requests.RequestException:
            if attempt < max_retries:
                print(f"\n{BColors.WARNING}[RETRYING]{BColors.ENDC} Network error for '{article['title']}' (Attempt {attempt + 2})")
    return None

# --- PHASE 4: DATABASE STORAGE ---
def store_analyzed_data(analyzed_data_list):
    """Stores analyzed article data in the database."""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    stored_count = 0
    for data in analyzed_data_list:
        try:
            summary_text = data.get('summary', 'N/A')
            recommendations_json = json.dumps(data.get('recommendations', []))
            cursor.execute("""
                INSERT OR IGNORE INTO articles
                (title, url, published_date, content, summary, threat_risk, category, recommendations)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                data['title'], data['url'], data['published_date'], data.get('content'),
                summary_text, data.get('threat_risk'), data.get('category'),
                recommendations_json
            ))
            if cursor.rowcount > 0:
                stored_count += 1
        except sqlite3.Error as e:
            log_error(f"DB error for {data['title']}: {e}")
    conn.commit()
    conn.close()
    log_success(f"Stored {stored_count} new articles in the database.")

# --- PHASE 5: REPORT GENERATION ---
def add_table_borders(table):
    """Adds borders to all cells in a table."""
    tbl_pr = table._element.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_elm = OxmlElement(f'w:{border_name}')
        border_elm.set(ns.qn('w:val'), 'single')
        border_elm.set(ns.qn('w:sz'), '4')
        border_elm.set(ns.qn('w:space'), '0')
        border_elm.set(ns.qn('w:color'), 'auto')
        tbl_borders.append(border_elm)
    tbl_pr.append(tbl_borders)

def set_cell_background_color(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), color_hex.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_bookmark(paragraph, bookmark_name):
    """Adds a bookmark to a paragraph."""
    run = paragraph.runs[0]
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(ns.qn('w:id'), '0')
    start.set(ns.qn('w:name'), bookmark_name)
    tag.addprevious(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(ns.qn('w:id'), '0')
    tag.addnext(end)

def add_internal_hyperlink(paragraph, text, anchor):
    """Adds an internal hyperlink to a paragraph."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(ns.qn('w:anchor'), anchor)
    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    style = OxmlElement('w:rStyle')
    style.set(ns.qn('w:val'), 'Hyperlink')
    run_props.append(style)
    run.append(run_props)
    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    for r in paragraph.runs:
        if r.text == text:
            r.font.name = 'Arial'
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.font.underline = True

def generate_weekly_report():
    """Generates the weekly threat intelligence report based on a flexible risk quota."""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    start_date, end_date = get_last_full_week_dates()
    cursor.execute("SELECT * FROM articles WHERE published_date BETWEEN ? AND ?", (start_date.isoformat(), end_date.isoformat()))
    all_weekly_articles = cursor.fetchall()
    conn.close()

    if not all_weekly_articles:
        log_warn("No articles found from the last week. No report generated.")
        return

    log_info(f"Found {len(all_weekly_articles)} total articles in DB for the week. Selecting based on risk quota...")
    categorized_articles = {"HIGH": [], "MEDIUM": [], "LOW": [], "INFORMATIONAL": []}
    for article in all_weekly_articles:
        risk = (article[6] or "N/A").upper()
        if risk in categorized_articles:
            categorized_articles[risk].append(article)

    report_quota = {"HIGH": 4, "MEDIUM": 3, "LOW": 2, "INFORMATIONAL": 1}
    target_report_size = 10
    weekly_data = []

    for risk_level in ["HIGH", "MEDIUM", "LOW", "INFORMATIONAL"]:
        count = report_quota[risk_level]
        available = categorized_articles.get(risk_level, [])
        selected = available[:count]
        weekly_data.extend(selected)
        for item in selected:
            categorized_articles[risk_level].remove(item)
        if len(selected) < count:
            log_warn(f"Requirement not met for {risk_level}: Found {len(selected)}, need {count}.")

    if len(weekly_data) < target_report_size:
        log_info(f"Report has {len(weekly_data)} articles, filling to {target_report_size} with best available.")
        remaining_pool = []
        for risk_level in ["HIGH", "MEDIUM", "LOW", "INFORMATIONAL"]:
            remaining_pool.extend(categorized_articles[risk_level])
        needed = target_report_size - len(weekly_data)
        weekly_data.extend(remaining_pool[:needed])

    risk_priority = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "INFORMATIONAL": 3}
    weekly_data.sort(key=lambda row: risk_priority.get((row[6] or "N/A").upper(), 99))

    if not weekly_data:
        log_warn("No articles selected for the report after applying quotas.")
        return
        
    log_info(f"Generating report with {len(weekly_data)} selected articles.")
    
    try:
        doc = Document(TEMPLATE_DOCX_PATH)
    except Exception as e:
        log_error(f"Could not open '{TEMPLATE_DOCX_PATH}'. Error: {e}")
        return

    # Document generation logic...
    period_p, titles_p, summary_intro_p = None, None, None
    for p in doc.paragraphs:
        if '[Placeholder-ReportingPeriod]' in p.text: period_p = p
        if '[Placeholder-Titles]' in p.text: titles_p = p
        if 'This report provides a detailed summary' in p.text: summary_intro_p = p

    if period_p:
        period_p.text = f"Reporting period: {start_date.strftime('%d %b %Y')} – {end_date.strftime('%d %b %Y')}"
        for run in period_p.runs: run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Arial'
        period_p._p.addnext(OxmlElement('w:p'))
    
    if titles_p:
        parent_body = titles_p._element.getparent()
        placeholder_index = parent_body.index(titles_p._element)
        for i, row_data in enumerate(weekly_data):
            p = doc.add_paragraph(f"{i+1}. ")
            add_internal_hyperlink(p, row_data[1], f"threat_{i+1}")
            p.paragraph_format.left_indent = Inches(0.5)
            p_element = p._element
            p_element.getparent().remove(p_element)
            parent_body.insert(placeholder_index + i, p_element)
        parent_body.remove(titles_p._element)

    if summary_intro_p: summary_intro_p._p.addnext(OxmlElement('w:p'))
    for table in doc.tables: table._tbl.getparent().remove(table._tbl)
    risk_colors = {"HIGH": "980000", "MEDIUM": "FFC000", "LOW": "38761D", "INFORMATIONAL": "1155CC", "N/A": "A9A9A9"}
    
    for i, article_data in enumerate(weekly_data):
        _id, title, url, _, _, summary_text, threat_risk, _, recommendations_json = article_data
        summary_text = summary_text or "Summary not available."
        threat_risk = (threat_risk or "N/A").upper()
        recommendations = json.loads(recommendations_json) if recommendations_json else []
        doc.add_paragraph()
        article_table = doc.add_table(rows=0, cols=1)
        add_table_borders(article_table)
        
        # Threat Risk Header
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"THREAT RISK = {threat_risk}"); run.font.name = 'Arial'; run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background_color(cell, risk_colors.get(threat_risk, "A9A9A9"))
        
        # Title
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title); run.font.name = 'Arial'; run.font.size = Pt(14); run.font.bold = True
        add_bookmark(p, f"threat_{i+1}")
        
        # Summary Section
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SUMMARY"); run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True
        set_cell_background_color(cell, "CCCCCC")
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; run = p.add_run(summary_text); run.font.name = 'Arial'; run.font.size = Pt(11)

        # Recommendations Section
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RECOMMENDATIONS"); run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True
        set_cell_background_color(cell, "CCCCCC")
        cell = article_table.add_row().cells[0]
        p_to_remove = cell.paragraphs[0]._element
        p_to_remove.getparent().remove(p_to_remove)

        for reco in recommendations:
            p = cell.add_paragraph()
            if isinstance(reco, dict) and 'title' in reco and 'description' in reco:
                p.add_run("• ").font.bold = True; p.add_run(f"{reco['title']}: ").font.bold = True; p.add_run(reco['description'])
                for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)
            else:
                reco_text = reco.get('step', reco) if isinstance(reco, dict) else reco
                run = p.add_run(f"• {reco_text}"); run.font.name = 'Arial'; run.font.size = Pt(11)
        
        # Sources Section
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Sources"); run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True
        set_cell_background_color(cell, "CCCCCC")
        cell = article_table.add_row().cells[0]
        p = cell.paragraphs[0]; run = p.add_run(f"• {url}"); run.font.name = 'Arial'; run.font.size = Pt(11)
    
    # Final Formatting
    for paragraph in doc.paragraphs:
        p_format = paragraph.paragraph_format; p_format.line_spacing = 1.15; p_format.space_before = Pt(6); p_format.space_after = Pt(6)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    p_format = paragraph.paragraph_format; p_format.line_spacing = 1.15; p_format.space_before = Pt(6); p_format.space_after = Pt(6)

    try:
        doc.save(OUTPUT_DOCX_PATH)
        log_success(f"Word report successfully generated and saved to: {OUTPUT_DOCX_PATH}")
    except Exception as e:
        log_error(f"Could not save the document. Error: {e}")

def main_pipeline():
    """Main function to run the full, sequential per-phase pipeline."""
    initialize_database()
    existing_urls = get_existing_urls()
    
    # Phase 1: Fetch and Scrape all new articles
    log_step(1, "Fetching and Scraping New Articles")
    new_articles = fetch_and_scrape_articles_sequential(existing_urls)

    # Phase 2: Filter for relevant articles
    log_step(2, "Filtering New Articles for Cybersecurity Relevance")
    relevant_articles = filter_articles_sequential(new_articles)

    if relevant_articles:
        # Phase 3: Analyze relevant articles
        log_step(3, "Analyzing New Relevant Articles with LLM")
        analyzed_data_list = analyze_articles_sequential(relevant_articles)
        
        if analyzed_data_list:
            # Phase 4: Store results in the database
            log_step(4, "Storing New Data in Database")
            store_analyzed_data(analyzed_data_list)
        else:
            log_warn("No new articles were successfully analyzed.")
    else:
        log_info("No new relevant articles to process.")
        
    # Phase 5: Generate the weekly report regardless of whether new articles were found
    log_step(5, "Generating Weekly Report")
    generate_weekly_report()
    log_success("Pipeline finished successfully!")


if __name__ == "__main__":
    DEBUG_MODE = "-debug" in sys.argv
    main_pipeline()

